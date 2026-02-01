"""DOCX document processing and text extraction"""
from docx import Document
from typing import Dict, List
import os


class DOCXProcessor:
    """Handles Word document processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = {
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'num_paragraphs': len(doc.paragraphs),
                'num_sections': len(doc.sections)
            }
            
            # Extract text
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = '\n\n'.join(paragraphs)
            
            return {
                'success': True,
                'full_text': full_text,
                'paragraphs': paragraphs,
                'metadata': metadata,
                'format': 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'docx'
            }
    
    def extract_with_structure(self, file_path: str) -> Dict[str, any]:
        """
        Extract text with structural information (headings, paragraphs, lists)
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with structured content
        """
        try:
            doc = Document(file_path)
            
            structured_content = []
            current_section = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    if current_section and current_section.get('content'):
                        structured_content.append(current_section)
                    
                    current_section = {
                        'heading': text,
                        'level': para.style.name,
                        'content': []
                    }
                else:
                    if current_section is None:
                        current_section = {
                            'heading': 'Introduction',
                            'level': 'Body',
                            'content': []
                        }
                    current_section['content'].append(text)
            
            if current_section and current_section.get('content'):
                structured_content.append(current_section)
            
            return {
                'success': True,
                'structured_content': structured_content,
                'format': 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'docx'
            }
